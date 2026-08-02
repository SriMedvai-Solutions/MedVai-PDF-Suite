"""Audit document and CSV generation for MedVai PDF Suite."""

from __future__ import annotations

import csv
import os
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import BatesQueueItem, NumberingConfig, PageRangeMap, RunSummary
from .utils import ensure_unique_filename


TECHNICAL_AUDIT_FOLDER = "Technical_Audit_Files"


class AuditGenerator:
    def __init__(self, output_folder: str):
        self.output_folder = output_folder
        self.technical_output_folder = os.path.join(output_folder, TECHNICAL_AUDIT_FOLDER)

    def _technical_path(self, filename: str) -> str:
        os.makedirs(self.technical_output_folder, exist_ok=True)
        return ensure_unique_filename(os.path.join(self.technical_output_folder, filename))

    def generate_audit(
        self,
        summary: RunSummary,
        bates_queue: Optional[list[BatesQueueItem]] = None,
        numbering_config: Optional[NumberingConfig] = None,
        preferred_base: Optional[str] = None,
    ) -> list[str]:
        if preferred_base:
            base = preferred_base
        elif summary.output_files:
            base = Path(summary.output_files[0]).stem
        else:
            base = "PDF_Processing_NEEDS_REVIEW"

        output_files: list[str] = []
        docx_path = ensure_unique_filename(os.path.join(self.output_folder, f"{base}_Audit.docx"))
        self._generate_docx(summary, bates_queue, numbering_config, docx_path)
        output_files.append(docx_path)

        if summary.page_mapping:
            path = self._technical_path(f"{base}_Page_Map.csv")
            self._generate_map_csv(summary.page_mapping, path)
            output_files.append(path)

        if bates_queue:
            path = self._technical_path(f"{base}_Bates_Map.csv")
            self._generate_bates_csv(bates_queue, path)
            output_files.append(path)

        stamp_audit = summary.settings.get("stamp_audit") or []
        if stamp_audit:
            path = self._technical_path(f"{base}_Stamp_Verification.csv")
            self._generate_stamp_csv(stamp_audit, path)
            output_files.append(path)

        return output_files

    def _generate_docx(
        self,
        summary: RunSummary,
        bates_queue: Optional[list[BatesQueueItem]],
        numbering_config: Optional[NumberingConfig],
        output_path: str,
    ) -> None:
        doc = Document()
        title = doc.add_heading("MedVai PDF Suite - Processing Audit", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Status: {summary.status}")
        doc.add_paragraph(f"Issue source: {summary.issue_source}")
        doc.add_paragraph(f"Operation: {summary.operation_type}")
        doc.add_paragraph(f"Version: {summary.settings.get('version', 'unknown')}")
        doc.add_paragraph(f"Build ID: {summary.settings.get('build_id', 'unknown')}")
        if summary.timestamp:
            doc.add_paragraph(f"Started: {summary.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        if summary.output_files:
            doc.add_paragraph(f"Main output: {Path(summary.output_files[0]).name}")
        else:
            doc.add_paragraph("Main output: No usable PDF was created.")

        doc.add_heading("1. Input and Output Summary", 1)
        doc.add_paragraph(f"Input files: {len(summary.input_files)}")
        for index, path in enumerate(summary.input_files, 1):
            doc.add_paragraph(f"{index}. {path}")
        doc.add_paragraph(f"Output pages: {summary.merged_pages}")

        section_number = 2
        if summary.page_mapping:
            doc.add_heading(f"{section_number}. Merge and Page Map", 1)
            section_number += 1
            self._add_page_map_table(doc, summary.page_mapping)

        if bates_queue:
            doc.add_heading(f"{section_number}. Bates Settings", 1)
            section_number += 1
            self._add_bates_table(doc, bates_queue)
            unselected = summary.settings.get("bates_unselected_sources") or []
            if unselected:
                doc.add_paragraph(
                    "The following source PDF occurrence(s) were merged without Bates because Bates was not selected:"
                )
                for item in unselected:
                    occurrence = int(item.get("occurrence", 1))
                    label = str(item.get("filename", ""))
                    if occurrence > 1:
                        label += f" (occurrence {occurrence})"
                    doc.add_paragraph(
                        f"• {label}: output pages {item.get('start_page', '')}–{item.get('end_page', '')}"
                    )

        if numbering_config:
            doc.add_heading(f"{section_number}. Numbering Settings", 1)
            section_number += 1
            self._add_numbering_table(doc, numbering_config)

        stamp_stats = summary.settings.get("stamp_stats") or {}
        stamp_audit = summary.settings.get("stamp_audit") or []
        if stamp_stats:
            doc.add_heading(f"{section_number}. Stamp Verification", 1)
            section_number += 1
            for key, label in (("bates", "Bates"), ("numbering", "Page numbering")):
                stats = stamp_stats.get(key, {})
                if stats.get("expected", 0) > 0:
                    doc.add_paragraph(
                        f"{label}: expected {stats.get('expected', 0)}, "
                        f"verified {stats.get('verified', 0)}, "
                        f"hard-stamped {stats.get('hard_stamp', 0)}, "
                        f"needs review {stats.get('needs_review', 0)}."
                    )

            review_entries = [entry for entry in stamp_audit if not entry.get("verified")]
            hard_entries = [entry for entry in stamp_audit if entry.get("mode") == "hard-stamp"]
            adjusted_entries = [entry for entry in stamp_audit if entry.get("adjustments")]

            if hard_entries:
                pages = sorted({int(entry["page"]) for entry in hard_entries})
                doc.add_paragraph(
                    "Hard-stamp fallback was used on output page(s): " + ", ".join(map(str, pages)) + "."
                )
                doc.add_paragraph(
                    "Those pages were visually flattened before the stamp was burned in. "
                    "Searchable text, links or annotations on those pages may no longer behave like the original."
                )

            if adjusted_entries:
                pages = sorted({int(entry["page"]) for entry in adjusted_entries})
                doc.add_paragraph(
                    "Placement was moved slightly inward or the font was reduced on output page(s): "
                    + ", ".join(map(str, pages)) + "."
                )

            if review_entries:
                doc.add_paragraph("Items that still need review:")
                table = doc.add_table(rows=1, cols=6)
                table.style = "Light Grid Accent 1"
                headings = ("Type", "Output page", "Source", "Expected text", "Mode", "Reason")
                for index, heading in enumerate(headings):
                    table.rows[0].cells[index].text = heading
                for entry in review_entries:
                    row = table.add_row().cells
                    row[0].text = str(entry.get("type", ""))
                    row[1].text = str(entry.get("page", ""))
                    source = str(entry.get("file", ""))
                    if source and entry.get("occurrence", 1) > 1:
                        source += f" (occurrence {entry['occurrence']})"
                    row[2].text = source
                    row[3].text = str(entry.get("text", ""))
                    row[4].text = str(entry.get("mode", ""))
                    row[5].text = str(entry.get("reason", ""))
            elif stamp_audit:
                doc.add_paragraph("Every expected stamp was found on the correct output page and inside its expected area.")

        if summary.warnings or summary.errors:
            doc.add_heading(f"{section_number}. Warnings and Items Needing Review", 1)
            if summary.warnings:
                doc.add_paragraph("Warnings:")
                for item in summary.warnings:
                    doc.add_paragraph(f"• {item}")
            if summary.errors:
                doc.add_paragraph("Items needing review:")
                for item in summary.errors:
                    doc.add_paragraph(f"• {item}")

        doc.add_paragraph("")
        doc.add_paragraph(
            "For legal or production use, review the main PDF together with this audit and the CSV verification files."
        )
        doc.save(output_path)

    @staticmethod
    def _add_page_map_table(doc: Document, page_mapping: list[PageRangeMap]) -> None:
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        headings = ("Order", "File", "Occurrence", "Source pages", "Output start", "Output end", "Entry ID")
        for index, heading in enumerate(headings):
            table.rows[0].cells[index].text = heading
        for order, mapping in enumerate(page_mapping, 1):
            row = table.add_row().cells
            row[0].text = str(order)
            row[1].text = mapping.filename
            row[2].text = str(mapping.occurrence)
            row[3].text = str(mapping.pages_in_source)
            row[4].text = str(mapping.start_page_in_merged)
            row[5].text = str(mapping.end_page_in_merged)
            row[6].text = mapping.entry_id

    @staticmethod
    def _add_bates_table(doc: Document, queue: list[BatesQueueItem]) -> None:
        table = doc.add_table(rows=1, cols=11)
        table.style = "Light Grid Accent 1"
        headings = (
            "Order", "File", "Prefix", "Symbol", "Start", "Padding", "Suffix",
            "Placement", "Font", "Colour", "First to last",
        )
        for index, heading in enumerate(headings):
            table.rows[0].cells[index].text = heading
        for order, item in enumerate(queue, 1):
            first = item.format_number(item.start_number)
            last = item.format_number(item.start_number + item.pages_in_source - 1)
            row = table.add_row().cells
            row[0].text = str(order)
            row[1].text = item.filename
            row[2].text = item.prefix
            row[3].text = item.symbol
            row[4].text = str(item.start_number)
            row[5].text = str(item.padding)
            row[6].text = item.suffix
            row[7].text = item.placement
            row[8].text = f"{item.font_name} {'Bold' if item.bold else ''} {item.font_size} pt".strip()
            row[9].text = item.color
            row[10].text = f"{first} to {last}"

    @staticmethod
    def _add_numbering_table(doc: Document, config: NumberingConfig) -> None:
        table = doc.add_table(rows=9, cols=2)
        table.style = "Light Grid Accent 1"
        values = (
            ("Pattern", config.pattern),
            ("Start number", str(config.start_number)),
            ("Placement", config.placement),
            ("Colour", config.color),
            ("Font", config.font_name),
            ("Bold", "Yes" if config.bold else "No"),
            ("Font size", str(config.font_size)),
            ("Opacity", str(config.opacity)),
            ("Offsets", f"X {config.offset_x}, Y {config.offset_y}"),
        )
        for index, (key, value) in enumerate(values):
            table.rows[index].cells[0].text = key
            table.rows[index].cells[1].text = value

    @staticmethod
    def _generate_map_csv(mapping: list[PageRangeMap], output_path: str) -> None:
        with open(output_path, "w", newline="", encoding="utf-8-sig") as handle:
            writer = csv.writer(handle)
            writer.writerow(("Order", "File", "Occurrence", "Source Pages", "Output Start", "Output End", "Entry ID", "Path"))
            for order, item in enumerate(mapping, 1):
                writer.writerow((
                    order, item.filename, item.occurrence, item.pages_in_source,
                    item.start_page_in_merged, item.end_page_in_merged, item.entry_id, item.file_path,
                ))

    @staticmethod
    def _generate_bates_csv(queue: list[BatesQueueItem], output_path: str) -> None:
        with open(output_path, "w", newline="", encoding="utf-8-sig") as handle:
            writer = csv.writer(handle)
            writer.writerow((
                "Order", "File", "Entry ID", "Prefix", "Symbol", "Start", "Padding", "Suffix",
                "Placement", "Colour", "Font", "Bold", "Font Size", "Opacity", "First Bates", "Last Bates",
            ))
            for order, item in enumerate(queue, 1):
                writer.writerow((
                    order, item.filename, item.entry_id, item.prefix, item.symbol, item.start_number,
                    item.padding, item.suffix, item.placement, item.color, item.font_name,
                    item.bold, item.font_size, item.opacity, item.format_number(item.start_number),
                    item.format_number(item.start_number + item.pages_in_source - 1),
                ))

    @staticmethod
    def _generate_stamp_csv(entries: list[dict[str, Any]], output_path: str) -> None:
        headings = (
            "Type", "Output Page", "Source File", "Occurrence", "Source Page", "Expected Text",
            "Verified", "Mode", "Reason", "Rotation", "Page Size", "Requested Placement",
            "Used Font Size", "Visual Rectangle", "Physical Rectangle", "Entry ID",
        )
        with open(output_path, "w", newline="", encoding="utf-8-sig") as handle:
            writer = csv.writer(handle)
            writer.writerow(headings)
            for entry in entries:
                writer.writerow((
                    entry.get("type", ""), entry.get("page", ""), entry.get("file", ""),
                    entry.get("occurrence", ""), entry.get("source_page", ""), entry.get("text", ""),
                    entry.get("verified", False), entry.get("mode", ""), entry.get("reason", ""),
                    entry.get("rotation", ""), entry.get("page_size", ""), entry.get("placement", ""),
                    entry.get("used_font_size", ""), entry.get("visual_rect", ""),
                    entry.get("physical_rect", ""), entry.get("entry_id", ""),
                ))
