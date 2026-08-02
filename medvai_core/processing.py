"""Core PDF processing for MedVai PDF Suite.

Safety design:
- input/folder-based output names;
- natural/manual order is preserved by the GUI and processing layer;
- duplicate source files are separate occurrences via entry_id;
- placement is calculated in the visually displayed page coordinate system;
- each expected stamp is verified after saving;
- difficult pages receive a direct retry and then a page-level hard-stamp fallback;
- outputs that cannot be fully verified are clearly named NEEDS_REVIEW.
"""

from __future__ import annotations

import os
import shutil
import tempfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Optional

import fitz
from pypdf import PdfReader, PdfWriter

from .audit import AuditGenerator, TECHNICAL_AUDIT_FOLDER
from .models import BatesQueueItem, NumberingConfig, PDFFileInfo, PageRangeMap, RunSummary
from .utils import (
    common_input_name,
    create_output_folder,
    ensure_unique_filename,
    hex_to_rgb_normalized,
    is_valid_hex_color,
    parse_page_ranges,
    safe_filename_component,
    validate_readable_pdf,
)


STATUS_VERIFIED = "Completed and Verified"
STATUS_REVIEW = "Completed — Needs Review"
STATUS_COULD_NOT_COMPLETE = "Could Not Complete — Needs Review"
BUILD_ID = "MEDVAI_PDF_SUITE_TECHNICAL_AUDIT_FOLDER_FIX_20260803"


class PDFProcessor:
    """Merge, split, Bates and page-number PDFs with verification."""

    def __init__(self, settings: Any):
        self.settings = settings
        self.warnings: list[str] = []
        self.errors: list[str] = []
        self.stamp_audit: list[dict[str, Any]] = []

    # ------------------------------------------------------------------
    # Public combined processing
    # ------------------------------------------------------------------
    def process_tab1(
        self,
        merge_enabled: bool,
        bates_enabled: bool,
        numbering_enabled: bool,
        pdf_files: list[PDFFileInfo],
        single_file: str,
        bates_queue: list[BatesQueueItem],
        numbering_config: Optional[NumberingConfig],
        output_folder: str,
        auto_subfolder: bool,
        output_base_name: str = "",
    ) -> dict[str, Any]:
        self.warnings = []
        self.errors = []
        self.stamp_audit = []
        started = datetime.now()
        operation_type = self._get_operation_type(merge_enabled, bates_enabled, numbering_enabled)

        if not any((merge_enabled, bates_enabled, numbering_enabled)):
            raise ValueError("Select at least one operation.")

        # The main PDF list always controls the merge order. Bates settings only
        # decide which merged source occurrences receive Bates. For backward
        # compatibility with older direct callers, an empty main list may still
        # be reconstructed from the Bates items.
        effective_pdf_files = list(pdf_files)
        if merge_enabled and not effective_pdf_files and bates_queue:
            effective_pdf_files = [
                PDFFileInfo(
                    path=item.file_path,
                    filename=item.filename,
                    pages=item.pages_in_source,
                    status="OK",
                    entry_id=item.source_entry_id or item.entry_id,
                )
                for item in bates_queue
            ]

        input_paths = [item.path for item in effective_pdf_files] if merge_enabled else [single_file]
        default_input_name = common_input_name(input_paths) if merge_enabled else safe_filename_component(Path(single_file).stem)
        input_name = safe_filename_component(output_base_name) if output_base_name.strip() else default_input_name
        final_output_folder = create_output_folder(output_folder, auto_subfolder, input_name)
        desired_filename = self._desired_output_name(
            input_name, merge_enabled, bates_enabled, numbering_enabled
        )
        run_stem, final_path, review_reserved_path = self._reserve_tab1_run(
            final_output_folder,
            Path(desired_filename).stem,
            bates_enabled=bates_enabled,
            stamping_enabled=bates_enabled or numbering_enabled,
        )

        page_mapping: list[PageRangeMap] = []
        total_pages = 0
        output_files: list[str] = []
        status = STATUS_REVIEW
        issue_source = "None identified"
        bates_unselected_sources: list[dict[str, Any]] = []

        with tempfile.TemporaryDirectory(prefix="medvai_pdf_", dir=final_output_folder) as temp_dir:
            try:
                self._validate_run_inputs(
                    merge_enabled=merge_enabled,
                    bates_enabled=bates_enabled,
                    numbering_enabled=numbering_enabled,
                    pdf_files=effective_pdf_files,
                    single_file=single_file,
                    bates_queue=bates_queue,
                    numbering_config=numbering_config,
                )

                current_input = single_file
                if merge_enabled:
                    merged_temp = os.path.join(temp_dir, "merged.pdf")
                    page_mapping = self._merge_pdfs(effective_pdf_files, merged_temp)
                    current_input = merged_temp

                if not merge_enabled:
                    source_pages = validate_readable_pdf(single_file)
                    page_mapping = [
                        PageRangeMap(
                            file_path=single_file,
                            filename=Path(single_file).name,
                            start_page_in_merged=1,
                            end_page_in_merged=source_pages,
                            pages_in_source=source_pages,
                            entry_id=(bates_queue[0].entry_id if bates_queue else "single"),
                            occurrence=1,
                        )
                    ]

                if bates_enabled:
                    selected_source_ids = {
                        item.source_entry_id or item.entry_id for item in bates_queue
                    }
                    if merge_enabled:
                        bates_unselected_sources = [
                            {
                                "filename": mapping.filename,
                                "occurrence": mapping.occurrence,
                                "start_page": mapping.start_page_in_merged,
                                "end_page": mapping.end_page_in_merged,
                            }
                            for mapping in page_mapping
                            if mapping.entry_id not in selected_source_ids
                        ]
                    bates_temp = os.path.join(temp_dir, "bates.pdf")
                    self._apply_perfile_bates(
                        current_input,
                        bates_queue,
                        page_mapping,
                        bates_temp,
                    )
                    current_input = bates_temp

                if numbering_enabled:
                    numbering_temp = os.path.join(temp_dir, "numbering.pdf")
                    assert numbering_config is not None
                    self._apply_global_numbering(current_input, numbering_config, numbering_temp)
                    current_input = numbering_temp

                # Merge-only comes here with merged.pdf as current_input.
                shutil.copy2(current_input, final_path)
                output_files = [final_path]
                total_pages = validate_readable_pdf(final_path)

                expected_pages = sum(item.pages for item in effective_pdf_files) if merge_enabled else validate_readable_pdf(single_file)
                if total_pages != expected_pages:
                    self.errors.append(
                        f"Output page count needs review: expected {expected_pages}, found {total_pages}."
                    )

                verification = {
                    "verified": True,
                    "hard_stamp_pages": [],
                    "unverified_entries": [],
                    "adjusted_entries": [],
                    "overlap_pages": [],
                }
                if bates_enabled or numbering_enabled:
                    verification = self._verify_and_repair_output(final_path)

                # A hard-stamped page is usable and verified, but must be disclosed.
                needs_review = bool(
                    self.errors
                    or verification["unverified_entries"]
                    or verification["hard_stamp_pages"]
                    or verification["adjusted_entries"]
                    or verification["overlap_pages"]
                )

                if verification["unverified_entries"]:
                    issue_source = "Placement or output verification issue"
                elif verification["hard_stamp_pages"]:
                    issue_source = "Difficult PDF page handled with hard-stamp fallback"
                elif verification["adjusted_entries"]:
                    issue_source = "Placement was automatically kept inside the visible page"
                elif verification["overlap_pages"]:
                    issue_source = "Bates and numbering overlap needs review"
                elif self.errors:
                    issue_source = "Output verification issue"

                status = STATUS_REVIEW if needs_review else STATUS_VERIFIED

                if status == STATUS_REVIEW:
                    os.replace(final_path, review_reserved_path)
                    final_path = review_reserved_path
                    output_files = [final_path]

            except Exception as exc:
                status = STATUS_COULD_NOT_COMPLETE
                issue_source = self._classify_issue(exc)
                self.errors.append(str(exc))
                if os.path.exists(final_path):
                    try:
                        os.remove(final_path)
                    except OSError:
                        pass
                output_files = []

        stamp_stats = self._compute_stamp_stats(
            bates_enabled=bates_enabled,
            numbering_enabled=numbering_enabled,
        )
        summary = RunSummary(
            timestamp=started,
            operation_type=operation_type,
            input_files=input_paths,
            output_files=output_files,
            output_path=output_files[0] if output_files else None,
            merged_pages=total_pages,
            bates_applied=stamp_stats["bates"]["verified"],
            warnings=self.warnings,
            errors=self.errors,
            success=status == STATUS_VERIFIED,
            status=status,
            issue_source=issue_source,
            page_mapping=page_mapping,
            settings={
                "merge": merge_enabled,
                "bates": bates_enabled,
                "numbering": numbering_enabled,
                "stamp_stats": stamp_stats,
                "stamp_audit": self.stamp_audit,
                "bates_unselected_sources": bates_unselected_sources,
                "version": "3.0.6-beta",
                "build_id": BUILD_ID,
            },
        )

        audit_generator = AuditGenerator(final_output_folder)
        audit_files = audit_generator.generate_audit(
            summary,
            bates_queue if bates_enabled else None,
            numbering_config if numbering_enabled else None,
            preferred_base=(Path(final_path).stem if output_files else run_stem),
        )
        log_path = self._write_log(
            final_output_folder,
            summary,
            preferred_base=(Path(final_path).stem if output_files else run_stem),
        )

        return {
            "success": status == STATUS_VERIFIED,
            "completed": bool(output_files),
            "status": status,
            "issue_source": issue_source,
            "output_folder": final_output_folder,
            "output_files": output_files,
            "audit_files": audit_files,
            "log_file": log_path,
            "stamp_stats": stamp_stats,
            "review_items": list(self.errors),
            "warnings": list(self.warnings),
        }

    # ------------------------------------------------------------------
    # Validation and naming
    # ------------------------------------------------------------------
    def _validate_run_inputs(
        self,
        *,
        merge_enabled: bool,
        bates_enabled: bool,
        numbering_enabled: bool,
        pdf_files: list[PDFFileInfo],
        single_file: str,
        bates_queue: list[BatesQueueItem],
        numbering_config: Optional[NumberingConfig],
    ) -> None:
        if merge_enabled:
            if not pdf_files:
                raise ValueError("Select at least one PDF to merge.")
            for info in pdf_files:
                pages = validate_readable_pdf(info.path)
                info.pages = pages
                info.status = "OK"
        else:
            validate_readable_pdf(single_file)

        if numbering_enabled:
            if numbering_config is None:
                raise ValueError("Numbering settings are missing.")
            self._validate_numbering_config(numbering_config)

        if bates_enabled:
            if not bates_queue:
                raise ValueError("Bates settings are missing.")
            for item in bates_queue:
                item.pages_in_source = validate_readable_pdf(item.file_path)
                self._validate_bates_item(item)

    def _validate_bates_item(self, item: BatesQueueItem) -> None:
        if item.start_number < 0:
            raise ValueError(f"Bates start number cannot be negative for {item.filename}.")
        if not (1 <= item.padding <= 20):
            raise ValueError("Bates number padding must be between 1 and 20 digits.")
        if item.prefix and (len(item.prefix) > 6 or not item.prefix.isalpha()):
            raise ValueError("Bates prefix may contain up to 6 letters only.")
        if len(item.suffix) > 50:
            raise ValueError("Bates suffix must be 50 characters or fewer.")
        if not is_valid_hex_color(item.color):
            raise ValueError(f"Invalid Bates colour for {item.filename}. Use #RRGGBB.")
        self._validate_stamp_style(
            item.placement,
            item.offset_x,
            item.offset_y,
            item.font_size,
            item.opacity,
            item.position_x_ratio,
            item.position_y_ratio,
        )

    def _validate_numbering_config(self, config: NumberingConfig) -> None:
        if "{n}" not in config.pattern:
            raise ValueError("The numbering pattern must contain {n}.")
        if config.start_number < 0:
            raise ValueError("The numbering start value cannot be negative.")
        if not is_valid_hex_color(config.color):
            raise ValueError("Invalid numbering colour. Use #RRGGBB.")
        self._validate_stamp_style(
            config.placement,
            config.offset_x,
            config.offset_y,
            config.font_size,
            config.opacity,
            config.position_x_ratio,
            config.position_y_ratio,
        )

    def _validate_stamp_style(
        self,
        placement: str,
        offset_x: int,
        offset_y: int,
        font_size: int,
        opacity: float,
        x_ratio: Optional[float],
        y_ratio: Optional[float],
    ) -> None:
        allowed = {
            "Top-Left", "Top-Center", "Top-Right",
            "Middle-Left", "Middle-Center", "Middle-Right",
            "Bottom-Left", "Bottom-Center", "Bottom-Right", "Custom",
        }
        if placement not in allowed:
            raise ValueError(f"Unknown placement: {placement}")
        if offset_x < 0 or offset_y < 0:
            raise ValueError("Position offsets cannot be negative.")
        if not (5 <= int(font_size) <= 72):
            raise ValueError("Font size must be between 5 and 72.")
        if not (0.1 <= float(opacity) <= 1.0):
            raise ValueError("Opacity must be between 0.1 and 1.0.")
        if placement == "Custom":
            if x_ratio is None or y_ratio is None:
                raise ValueError("Open Preview & Set Placement and save the custom position.")
            if not (0.0 <= x_ratio <= 1.0 and 0.0 <= y_ratio <= 1.0):
                raise ValueError("The saved custom position is outside the visible page.")

    def _desired_output_name(self, base: str, merge: bool, bates: bool, numbering: bool) -> str:
        base = safe_filename_component(base)
        if merge:
            suffix = "_Merged"
            if bates:
                suffix += "_Bates"
            if numbering:
                suffix += "_Numbered"
            return f"{base}{suffix}.pdf"
        if bates and numbering:
            return f"{base}_Bates_Numbered.pdf"
        if bates:
            return f"{base}_Bates.pdf"
        return f"{base}_Numbered.pdf"


    @staticmethod
    def _reserve_tab1_run(
        output_folder: str,
        desired_stem: str,
        *,
        bates_enabled: bool,
        stamping_enabled: bool,
    ) -> tuple[str, str, str]:
        """Reserve one numeric run suffix across the main and technical folders."""
        technical_folder = Path(output_folder) / TECHNICAL_AUDIT_FOLDER
        counter = 0
        while True:
            suffix = "" if counter == 0 else f"_{counter}"
            run_stem = f"{desired_stem}{suffix}"
            review_stem = f"{run_stem}_NEEDS_REVIEW"
            candidates = [
                Path(output_folder) / f"{run_stem}.pdf",
                Path(output_folder) / f"{review_stem}.pdf",
            ]
            for base in (run_stem, review_stem):
                candidates.append(Path(output_folder) / f"{base}_Audit.docx")
                candidates.extend([
                    technical_folder / f"{base}_Page_Map.csv",
                    technical_folder / f"{base}_Process_Log.txt",
                ])
                if bates_enabled:
                    candidates.append(technical_folder / f"{base}_Bates_Map.csv")
                if stamping_enabled:
                    candidates.append(technical_folder / f"{base}_Stamp_Verification.csv")
            if not any(candidate.exists() for candidate in candidates):
                return (
                    run_stem,
                    str(Path(output_folder) / f"{run_stem}.pdf"),
                    str(Path(output_folder) / f"{review_stem}.pdf"),
                )
            counter += 1

    @staticmethod
    def _needs_review_path(path: str) -> str:
        source = Path(path)
        return ensure_unique_filename(str(source.with_name(f"{source.stem}_NEEDS_REVIEW{source.suffix}")))

    @staticmethod
    def _get_operation_type(merge: bool, bates: bool, numbering: bool) -> str:
        parts = []
        if merge:
            parts.append("Merge")
        if bates:
            parts.append("Bates")
        if numbering:
            parts.append("Numbering")
        return " + ".join(parts)

    @staticmethod
    def _classify_issue(exc: Exception) -> str:
        text = str(exc).casefold()
        if any(word in text for word in ("password", "unlock", "damaged", "could not be opened", "no pages")):
            return "Input PDF issue"
        if any(word in text for word in ("placement", "offset", "font", "colour", "color", "position", "pattern")):
            return "Placement or settings issue"
        if "output folder" in text or "writable" in text or "storage" in text:
            return "Computer or output-folder issue"
        return "Tool processing issue"

    # ------------------------------------------------------------------
    # Merge
    # ------------------------------------------------------------------
    def _merge_pdfs(self, pdf_files: list[PDFFileInfo], output_path: str) -> list[PageRangeMap]:
        writer = PdfWriter()
        page_mapping: list[PageRangeMap] = []
        current_page = 1
        occurrences: Counter[str] = Counter()
        first_metadata: Optional[dict[str, str]] = None

        for info in pdf_files:
            reader = PdfReader(info.path)
            if reader.is_encrypted:
                raise ValueError(
                    f"This PDF is password-protected: {Path(info.path).name}. Please unlock it and try again."
                )
            total_pages = len(reader.pages)
            resolved = str(Path(info.path).resolve()).casefold()
            occurrences[resolved] += 1
            occurrence = occurrences[resolved]
            label = Path(info.path).name
            if occurrence > 1:
                label = f"{label} — Occurrence {occurrence}"

            page_mapping.append(
                PageRangeMap(
                    file_path=info.path,
                    filename=Path(info.path).name,
                    start_page_in_merged=current_page,
                    end_page_in_merged=current_page + total_pages - 1,
                    pages_in_source=total_pages,
                    entry_id=info.entry_id,
                    occurrence=occurrence,
                )
            )

            # append preserves page objects, links and imported outlines, while
            # adding a clear top-level bookmark for every input occurrence.
            writer.append(reader, outline_item=label, import_outline=True)
            if first_metadata is None and reader.metadata:
                first_metadata = {
                    str(key): str(value)
                    for key, value in reader.metadata.items()
                    if key and value is not None
                }
            current_page += total_pages

        if first_metadata:
            try:
                writer.add_metadata(first_metadata)
            except Exception as exc:
                self.warnings.append(f"Source metadata could not be fully preserved: {exc}")

        with open(output_path, "wb") as file_handle:
            writer.write(file_handle)
        return page_mapping

    # ------------------------------------------------------------------
    # Stamp placement and insertion
    # ------------------------------------------------------------------
    @staticmethod
    def _fitz_font_name(font_name: str, bold: bool) -> str:
        name = (font_name or "Helvetica").casefold()
        if "times" in name:
            return "tibo" if bold else "tiro"
        if "courier" in name:
            return "cobo" if bold else "cour"
        return "hebo" if bold else "helv"

    def _calculate_visual_rect(
        self,
        page: fitz.Page,
        text: str,
        *,
        placement: str,
        font_name: str,
        bold: bool,
        font_size: int,
        offset_x: int,
        offset_y: int,
        position_x_ratio: Optional[float],
        position_y_ratio: Optional[float],
    ) -> tuple[fitz.Rect, int, list[str]]:
        page_width = float(page.rect.width)
        page_height = float(page.rect.height)
        safe = 4.0
        font_code = self._fitz_font_name(font_name, bold)
        used_font_size = int(font_size)
        text_width = fitz.get_text_length(text, fontname=font_code, fontsize=used_font_size)
        text_height = used_font_size * 1.25
        adjustments: list[str] = []

        available_width = page_width - (safe * 2)
        if text_width > available_width:
            scale = available_width / max(text_width, 1.0)
            reduced = max(5, int(used_font_size * scale * 0.98))
            if reduced >= used_font_size or reduced < 5:
                raise ValueError(
                    f"The stamp text is too wide for page {page.number + 1}, even at the minimum font size."
                )
            used_font_size = reduced
            text_width = fitz.get_text_length(text, fontname=font_code, fontsize=used_font_size)
            text_height = used_font_size * 1.25
            adjustments.append(f"font reduced to {used_font_size} to remain visible")

        if placement == "Custom":
            assert position_x_ratio is not None and position_y_ratio is not None
            center_x = position_x_ratio * page_width
            center_y = position_y_ratio * page_height
            x0 = center_x - text_width / 2
            y0 = center_y - text_height / 2
        else:
            if "Left" in placement:
                if offset_x + text_width + safe > page_width:
                    raise ValueError(
                        f"The X offset places the stamp outside page {page.number + 1}."
                    )
                x0 = float(offset_x)
            elif "Right" in placement:
                if offset_x + text_width + safe > page_width:
                    raise ValueError(
                        f"The X offset places the stamp outside page {page.number + 1}."
                    )
                x0 = page_width - float(offset_x) - text_width
            else:
                x0 = (page_width - text_width) / 2

            if "Top" in placement:
                if offset_y + text_height + safe > page_height:
                    raise ValueError(
                        f"The Y offset places the stamp outside page {page.number + 1}."
                    )
                y0 = float(offset_y)
            elif "Bottom" in placement:
                if offset_y + text_height + safe > page_height:
                    raise ValueError(
                        f"The Y offset places the stamp outside page {page.number + 1}."
                    )
                y0 = page_height - float(offset_y) - text_height
            else:
                y0 = (page_height - text_height) / 2

        # Custom ratio stays visually consistent. For an unusually small page,
        # move only the minimum distance needed to keep the full text visible.
        clamped_x0 = min(max(x0, safe), page_width - safe - text_width)
        clamped_y0 = min(max(y0, safe), page_height - safe - text_height)
        if abs(clamped_x0 - x0) > 0.25 or abs(clamped_y0 - y0) > 0.25:
            adjustments.append("position moved slightly inward to keep the full stamp visible")
        x0, y0 = clamped_x0, clamped_y0

        rect = fitz.Rect(x0, y0, x0 + text_width, y0 + text_height)
        if rect.x0 < 0 or rect.y0 < 0 or rect.x1 > page_width or rect.y1 > page_height:
            raise ValueError(f"The stamp would be outside the visible area on page {page.number + 1}.")
        return rect, used_font_size, adjustments

    @staticmethod
    def _visual_rect_to_physical(page: fitz.Page, visual_rect: fitz.Rect) -> fitz.Rect:
        translated = fitz.Rect(
            page.rect.x0 + visual_rect.x0,
            page.rect.y0 + visual_rect.y0,
            page.rect.x0 + visual_rect.x1,
            page.rect.y0 + visual_rect.y1,
        )
        if page.rotation == 0:
            return translated
        points = [
            fitz.Point(translated.x0, translated.y0) * page.derotation_matrix,
            fitz.Point(translated.x1, translated.y0) * page.derotation_matrix,
            fitz.Point(translated.x0, translated.y1) * page.derotation_matrix,
            fitz.Point(translated.x1, translated.y1) * page.derotation_matrix,
        ]
        return fitz.Rect(
            min(point.x for point in points),
            min(point.y for point in points),
            max(point.x for point in points),
            max(point.y for point in points),
        )

    def _insert_stamp(
        self,
        page: fitz.Page,
        text: str,
        *,
        placement: str,
        color: str,
        font_name: str,
        bold: bool,
        font_size: int,
        opacity: float,
        offset_x: int,
        offset_y: int,
        position_x_ratio: Optional[float],
        position_y_ratio: Optional[float],
    ) -> dict[str, Any]:
        visual_rect, used_font_size, adjustments = self._calculate_visual_rect(
            page,
            text,
            placement=placement,
            font_name=font_name,
            bold=bold,
            font_size=font_size,
            offset_x=offset_x,
            offset_y=offset_y,
            position_x_ratio=position_x_ratio,
            position_y_ratio=position_y_ratio,
        )
        font_code = self._fitz_font_name(font_name, bold)
        rgb = hex_to_rgb_normalized(color)

        visual_baseline = fitz.Point(
            page.rect.x0 + visual_rect.x0,
            page.rect.y0 + visual_rect.y0 + used_font_size,
        )
        if page.rotation:
            point = visual_baseline * page.derotation_matrix
            text_rotation = page.rotation
            text_width = fitz.get_text_length(text, fontname=font_code, fontsize=used_font_size)
            if page.rotation == 180:
                point = fitz.Point(point.x + text_width, point.y - used_font_size)
        else:
            point = visual_baseline
            text_rotation = 0

        page.insert_text(
            point,
            text,
            fontname=font_code,
            fontsize=used_font_size,
            color=rgb,
            rotate=text_rotation,
            overlay=True,
            fill_opacity=float(opacity),
            stroke_opacity=float(opacity),
        )
        physical_rect = self._visual_rect_to_physical(page, visual_rect)
        return {
            "visual_rect": [visual_rect.x0, visual_rect.y0, visual_rect.x1, visual_rect.y1],
            "physical_rect": [physical_rect.x0, physical_rect.y0, physical_rect.x1, physical_rect.y1],
            "used_font_size": used_font_size,
            "adjustments": adjustments,
        }

    def _apply_perfile_bates(
        self,
        input_pdf: str,
        bates_queue: list[BatesQueueItem],
        page_mapping: list[PageRangeMap],
        output_path: str,
    ) -> None:
        bates_by_source: dict[str, BatesQueueItem] = {}
        for item in bates_queue:
            key = item.source_entry_id or item.entry_id
            bates_by_source.setdefault(key, item)
        with fitz.open(input_pdf) as doc:
            for mapping in page_mapping:
                item = bates_by_source.get(mapping.entry_id)
                if item is None:
                    # A merged source without Bates settings is intentionally
                    # retained and left unstamped. For a single-file Bates run,
                    # keep backward compatibility with one Bates item.
                    if len(page_mapping) == 1 and len(bates_queue) == 1:
                        item = bates_queue[0]
                    else:
                        continue

                for output_page in range(mapping.start_page_in_merged, mapping.end_page_in_merged + 1):
                    source_page = output_page - mapping.start_page_in_merged + 1
                    bates_number = item.start_number + source_page - 1
                    text = item.format_number(bates_number)
                    page = doc[output_page - 1]
                    entry = self._new_stamp_entry(
                        stamp_type="bates",
                        page=page,
                        output_page=output_page,
                        source_page=source_page,
                        filename=mapping.filename,
                        occurrence=mapping.occurrence,
                        entry_id=item.entry_id,
                        text=text,
                        placement=item.placement,
                        config=self._bates_config_dict(item),
                    )
                    try:
                        result = self._insert_stamp(page, text, **entry["config"])
                        entry.update(result)
                        if result["adjustments"]:
                            entry["mode"] = "auto-adjusted"
                            entry["reason"] = "; ".join(result["adjustments"])
                    except Exception as exc:
                        entry["mode"] = "needs-review"
                        entry["reason"] = str(exc)
                        self.errors.append(
                            f"Bates placement needs review on output page {output_page}: {exc}"
                        )
                    self.stamp_audit.append(entry)
            doc.save(output_path, garbage=4, deflate=True)

    def _apply_global_numbering(
        self,
        input_pdf: str,
        config: NumberingConfig,
        output_path: str,
    ) -> None:
        with fitz.open(input_pdf) as doc:
            total_pages = len(doc)
            for index, page in enumerate(doc):
                output_page = index + 1
                display_number = config.start_number + index + config.adjust
                text = config.pattern.replace("{n}", str(display_number)).replace("{total}", str(total_pages))
                entry = self._new_stamp_entry(
                    stamp_type="numbering",
                    page=page,
                    output_page=output_page,
                    source_page=output_page,
                    filename="",
                    occurrence=1,
                    entry_id="global-numbering",
                    text=text,
                    placement=config.placement,
                    config=self._numbering_config_dict(config),
                )
                try:
                    result = self._insert_stamp(page, text, **entry["config"])
                    entry.update(result)
                    if result["adjustments"]:
                        entry["mode"] = "auto-adjusted"
                        entry["reason"] = "; ".join(result["adjustments"])
                except Exception as exc:
                    entry["mode"] = "needs-review"
                    entry["reason"] = str(exc)
                    self.errors.append(
                        f"Page numbering placement needs review on output page {output_page}: {exc}"
                    )
                self.stamp_audit.append(entry)
            doc.save(output_path, garbage=4, deflate=True)

    @staticmethod
    def _bates_config_dict(item: BatesQueueItem) -> dict[str, Any]:
        return {
            "placement": item.placement,
            "color": item.color,
            "font_name": item.font_name,
            "bold": item.bold,
            "font_size": item.font_size,
            "opacity": item.opacity,
            "offset_x": item.offset_x,
            "offset_y": item.offset_y,
            "position_x_ratio": item.position_x_ratio,
            "position_y_ratio": item.position_y_ratio,
        }

    @staticmethod
    def _numbering_config_dict(config: NumberingConfig) -> dict[str, Any]:
        return {
            "placement": config.placement,
            "color": config.color,
            "font_name": config.font_name,
            "bold": config.bold,
            "font_size": config.font_size,
            "opacity": config.opacity,
            "offset_x": config.offset_x,
            "offset_y": config.offset_y,
            "position_x_ratio": config.position_x_ratio,
            "position_y_ratio": config.position_y_ratio,
        }

    @staticmethod
    def _new_stamp_entry(
        *,
        stamp_type: str,
        page: fitz.Page,
        output_page: int,
        source_page: int,
        filename: str,
        occurrence: int,
        entry_id: str,
        text: str,
        placement: str,
        config: dict[str, Any],
    ) -> dict[str, Any]:
        return {
            "type": stamp_type,
            "file": filename,
            "occurrence": occurrence,
            "source_page": source_page,
            "page": output_page,
            "entry_id": entry_id,
            "text": text,
            "placement": placement,
            "mode": "normal",
            "reason": "",
            "verified": False,
            "rotation": int(page.rotation),
            "page_size": f"{page.rect.width:.2f} x {page.rect.height:.2f}",
            "visual_rect": None,
            "physical_rect": None,
            "used_font_size": config["font_size"],
            "adjustments": [],
            "config": dict(config),
        }

    # ------------------------------------------------------------------
    # Verification, retry and hard-stamp fallback
    # ------------------------------------------------------------------
    @staticmethod
    def _rects_intersect(left: fitz.Rect, right: fitz.Rect, tolerance: float = 0.0) -> bool:
        a = fitz.Rect(left.x0 - tolerance, left.y0 - tolerance, left.x1 + tolerance, left.y1 + tolerance)
        return not (a.x1 < right.x0 or a.x0 > right.x1 or a.y1 < right.y0 or a.y0 > right.y1)

    def _verify_entry(self, page: fitz.Page, entry: dict[str, Any]) -> bool:
        if not entry.get("physical_rect"):
            return False
        expected = fitz.Rect(entry["physical_rect"])
        try:
            hits = page.search_for(entry["text"])
        except Exception:
            return False
        return any(self._rects_intersect(expected, hit, tolerance=8.0) for hit in hits)

    def _verify_all_entries(self, pdf_path: str) -> list[dict[str, Any]]:
        missing: list[dict[str, Any]] = []
        with fitz.open(pdf_path) as doc:
            for entry in self.stamp_audit:
                page_index = int(entry["page"]) - 1
                if page_index < 0 or page_index >= len(doc):
                    entry["verified"] = False
                    entry["reason"] = "Output page does not exist."
                    missing.append(entry)
                    continue
                verified = self._verify_entry(doc[page_index], entry)
                entry["verified"] = verified
                if not verified:
                    missing.append(entry)
        return missing

    def _retry_missing_entries(self, pdf_path: str, missing: list[dict[str, Any]]) -> None:
        if not missing:
            return
        retry_path = f"{pdf_path}.retry.pdf"
        with fitz.open(pdf_path) as doc:
            for entry in missing:
                page = doc[int(entry["page"]) - 1]
                try:
                    result = self._insert_stamp(page, entry["text"], **entry["config"])
                    entry.update(result)
                    entry["mode"] = "retry"
                    entry["reason"] = "Direct insertion was retried with rotation-aware coordinates."
                except Exception as exc:
                    entry["mode"] = "needs-review"
                    entry["reason"] = f"Direct retry could not be applied: {exc}"
            doc.save(retry_path, garbage=4, deflate=True)
        os.replace(retry_path, pdf_path)

    def _hard_stamp_pages(self, pdf_path: str, page_numbers: Iterable[int]) -> list[int]:
        pages = sorted(set(int(page) for page in page_numbers))
        if not pages:
            return []
        hard_path = f"{pdf_path}.hard.pdf"
        entries_by_page: dict[int, list[dict[str, Any]]] = defaultdict(list)
        for entry in self.stamp_audit:
            entries_by_page[int(entry["page"])].append(entry)

        completed: list[int] = []
        with fitz.open(pdf_path) as doc:
            for page_number in pages:
                index = page_number - 1
                if index < 0 or index >= len(doc):
                    continue
                old_page = doc[index]
                try:
                    pixmap = old_page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False, annots=True)
                    image_bytes = pixmap.tobytes("png")
                    width, height = float(old_page.rect.width), float(old_page.rect.height)
                    new_page = doc.new_page(pno=index, width=width, height=height)
                    new_page.insert_image(new_page.rect, stream=image_bytes, overlay=True)

                    # Reinsert every expected stamp before deleting the old page.
                    # A structural page deletion invalidates existing Page objects.
                    for entry in entries_by_page[page_number]:
                        result = self._insert_stamp(new_page, entry["text"], **entry["config"])
                        entry.update(result)
                        entry["mode"] = "hard-stamp"
                        entry["reason"] = (
                            "The page was visually flattened and the stamp was burned into the replacement page "
                            "because normal verification did not pass. Searchable text, links or annotations on this "
                            "page may no longer behave like the original."
                        )
                        entry["rotation"] = 0
                        entry["page_size"] = f"{new_page.rect.width:.2f} x {new_page.rect.height:.2f}"
                    doc.delete_page(index + 1)
                    completed.append(page_number)
                except Exception as exc:
                    self.errors.append(
                        f"Output page {page_number} could not receive the hard-stamp fallback: {exc}"
                    )
            doc.save(hard_path, garbage=4, deflate=True)
        os.replace(hard_path, pdf_path)
        return completed

    def _find_overlap_pages(self) -> list[int]:
        by_page: dict[int, list[dict[str, Any]]] = defaultdict(list)
        for entry in self.stamp_audit:
            if entry.get("physical_rect"):
                by_page[int(entry["page"])].append(entry)
        overlaps: list[int] = []
        for page, entries in by_page.items():
            bates = [entry for entry in entries if entry["type"] == "bates"]
            numbering = [entry for entry in entries if entry["type"] == "numbering"]
            if any(
                self._rects_intersect(fitz.Rect(b["physical_rect"]), fitz.Rect(n["physical_rect"]), tolerance=1.0)
                for b in bates for n in numbering
            ):
                overlaps.append(page)
        return sorted(overlaps)

    def _verify_and_repair_output(self, pdf_path: str) -> dict[str, Any]:
        missing = self._verify_all_entries(pdf_path)
        if missing:
            self._retry_missing_entries(pdf_path, missing)
            missing = self._verify_all_entries(pdf_path)

        hard_pages: list[int] = []
        if missing:
            hard_pages = self._hard_stamp_pages(pdf_path, [entry["page"] for entry in missing])
            missing = self._verify_all_entries(pdf_path)

        for entry in missing:
            entry["mode"] = "needs-review"
            if not entry.get("reason"):
                entry["reason"] = "The expected stamp could not be verified in the saved output."
            self.errors.append(
                f"{entry['type'].title()} needs review on output page {entry['page']}: {entry['text']}"
            )

        adjusted = [entry for entry in self.stamp_audit if entry.get("adjustments")]
        overlaps = self._find_overlap_pages()
        if overlaps:
            self.warnings.append(
                "Bates and page numbering overlap on output page(s): " + ", ".join(map(str, overlaps))
            )

        return {
            "verified": not missing,
            "hard_stamp_pages": hard_pages,
            "unverified_entries": missing,
            "adjusted_entries": adjusted,
            "overlap_pages": overlaps,
        }

    def _compute_stamp_stats(self, *, bates_enabled: bool, numbering_enabled: bool) -> dict[str, dict[str, int]]:
        stats = {
            "bates": {"expected": 0, "verified": 0, "hard_stamp": 0, "needs_review": 0},
            "numbering": {"expected": 0, "verified": 0, "hard_stamp": 0, "needs_review": 0},
        }
        for entry in self.stamp_audit:
            key = "bates" if entry["type"] == "bates" else "numbering"
            stats[key]["expected"] += 1
            if entry.get("verified"):
                stats[key]["verified"] += 1
            else:
                stats[key]["needs_review"] += 1
            if entry.get("mode") == "hard-stamp":
                stats[key]["hard_stamp"] += 1
        if not bates_enabled:
            stats["bates"] = {"expected": 0, "verified": 0, "hard_stamp": 0, "needs_review": 0}
        if not numbering_enabled:
            stats["numbering"] = {"expected": 0, "verified": 0, "hard_stamp": 0, "needs_review": 0}
        return stats

    # ------------------------------------------------------------------
    # Split
    # ------------------------------------------------------------------
    def split_by_ranges(
        self,
        input_file: str,
        output_folder: str,
        ranges_str: str,
        auto_subfolder: bool = True,
    ) -> dict[str, Any]:
        page_count = validate_readable_pdf(input_file)
        ranges = parse_page_ranges(ranges_str, page_count=page_count)
        input_name = safe_filename_component(Path(input_file).stem)
        final_output_folder = create_output_folder(output_folder, auto_subfolder, input_name)

        pdf_paths, summary_path, log_path = self._reserve_split_run(
            final_output_folder, input_name, ranges
        )

        reader = PdfReader(input_file)
        actual_ranges: list[tuple[int, int, str]] = []
        moved_paths: list[str] = []
        with tempfile.TemporaryDirectory(prefix="medvai_split_", dir=final_output_folder) as temp_dir:
            temp_pdf_paths: list[str] = []
            for index, ((start, end), output_path) in enumerate(zip(ranges, pdf_paths), 1):
                writer = PdfWriter()
                for page_number in range(start - 1, end):
                    writer.add_page(reader.pages[page_number])
                if len(writer.pages) != end - start + 1:
                    raise RuntimeError(f"The requested range {start}-{end} could not be created completely.")
                temp_path = str(Path(temp_dir) / f"part_{index}.pdf")
                with open(temp_path, "wb") as handle:
                    writer.write(handle)
                actual_count = validate_readable_pdf(temp_path)
                if actual_count != end - start + 1:
                    raise RuntimeError(
                        f"The output {Path(output_path).name} needs review because its page count is incorrect."
                    )
                temp_pdf_paths.append(temp_path)
                actual_ranges.append((start, end, output_path))

            temp_summary = str(Path(temp_dir) / "summary.docx")
            temp_log = str(Path(temp_dir) / "process_log.txt")
            self._generate_split_summary(input_file, actual_ranges, temp_summary)
            self._write_split_log(input_file, actual_ranges, temp_log)

            try:
                for temp_path, output_path in zip(temp_pdf_paths, pdf_paths):
                    os.replace(temp_path, output_path)
                    moved_paths.append(output_path)
                os.replace(temp_summary, summary_path)
                moved_paths.append(summary_path)
                os.makedirs(str(Path(log_path).parent), exist_ok=True)
                os.replace(temp_log, log_path)
                moved_paths.append(log_path)
            except Exception:
                for moved in moved_paths:
                    try:
                        os.remove(moved)
                    except OSError:
                        pass
                raise

        return {
            "success": True,
            "completed": True,
            "status": STATUS_VERIFIED,
            "issue_source": "None identified",
            "output_folder": final_output_folder,
            "output_files": pdf_paths,
            "audit_files": [summary_path],
            "log_file": log_path,
            "stamp_stats": {},
            "review_items": [],
            "warnings": [],
        }

    @staticmethod
    def _reserve_split_run(
        output_folder: str,
        input_name: str,
        ranges: list[tuple[int, int]],
    ) -> tuple[list[str], str, str]:
        """Reserve one shared suffix for split PDFs, summary and technical log."""
        technical_folder = Path(output_folder) / TECHNICAL_AUDIT_FOLDER
        counter = 0
        while True:
            suffix = "" if counter == 0 else f"_{counter}"
            pdf_paths = [
                str(Path(output_folder) / f"{input_name}_Split_{start}-{end}{suffix}.pdf")
                for start, end in ranges
            ]
            summary_path = str(
                Path(output_folder) / f"{input_name}_Split_Summary{suffix}.docx"
            )
            log_path = str(
                technical_folder / f"{input_name}_Split_Process_Log{suffix}.txt"
            )
            candidates = [*pdf_paths, summary_path, log_path]
            if not any(Path(candidate).exists() for candidate in candidates):
                return pdf_paths, summary_path, log_path
            counter += 1

    def split_by_pages(
        self,
        input_file: str,
        output_folder: str,
        pages_per_file: int,
        auto_subfolder: bool = True,
    ) -> dict[str, Any]:
        page_count = validate_readable_pdf(input_file)
        if not isinstance(pages_per_file, int) or pages_per_file < 1:
            raise ValueError("Pages per file must be a whole number of at least 1.")
        ranges = [
            (start, min(start + pages_per_file - 1, page_count))
            for start in range(1, page_count + 1, pages_per_file)
        ]
        range_text = ", ".join(f"{start}-{end}" for start, end in ranges)
        return self.split_by_ranges(input_file, output_folder, range_text, auto_subfolder)

    @staticmethod
    def _generate_split_summary(
        input_file: str,
        outputs: list[tuple[int, int, str]],
        summary_path: str,
    ) -> None:
        from docx import Document

        doc = Document()
        doc.add_heading("PDF Split Summary", 0)
        doc.add_paragraph(f"Status: {STATUS_VERIFIED}")
        doc.add_paragraph(f"Input: {Path(input_file).name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total output files: {len(outputs)}")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for index, heading in enumerate(("Part", "Filename", "Actual pages", "Page count")):
            table.rows[0].cells[index].text = heading
        for part, (start, end, output_path) in enumerate(outputs, 1):
            row = table.add_row().cells
            row[0].text = str(part)
            row[1].text = Path(output_path).name
            row[2].text = f"{start}-{end}"
            row[3].text = str(end - start + 1)
        doc.save(summary_path)

    @staticmethod
    def _write_split_log(
        input_file: str,
        ranges: list[tuple[int, int, str]],
        output_path: str,
    ) -> None:
        with open(output_path, "w", encoding="utf-8") as handle:
            handle.write("MedVai PDF Suite Split Process Log\n")
            handle.write("=" * 72 + "\n")
            handle.write(f"Version: 3.0.6-beta\n")
            handle.write(f"Build ID: {BUILD_ID}\n")
            handle.write(f"Input: {input_file}\n")
            handle.write(f"Status: {STATUS_VERIFIED}\n")
            handle.write(f"Outputs: {len(ranges)}\n")
            for start, end, path in ranges:
                handle.write(
                    f"  - {Path(path).name}: source pages {start}-{end} "
                    f"({end - start + 1} pages)\n"
                )

    # ------------------------------------------------------------------
    # Log
    # ------------------------------------------------------------------
    def _write_log(self, output_folder: str, summary: RunSummary, preferred_base: str) -> str:
        technical_folder = os.path.join(output_folder, TECHNICAL_AUDIT_FOLDER)
        os.makedirs(technical_folder, exist_ok=True)
        log_path = ensure_unique_filename(
            os.path.join(technical_folder, f"{preferred_base}_Process_Log.txt")
        )
        with open(log_path, "w", encoding="utf-8") as handle:
            handle.write("MedVai PDF Suite Process Log\n")
            handle.write("=" * 72 + "\n")
            handle.write(f"Version: {summary.settings.get('version', 'unknown')}\n")
            handle.write(f"Build ID: {summary.settings.get('build_id', 'unknown')}\n")
            handle.write(f"Started: {summary.timestamp.strftime('%Y-%m-%d %H:%M:%S')}\n")
            handle.write(f"Operation: {summary.operation_type}\n")
            handle.write(f"Status: {summary.status}\n")
            handle.write(f"Issue source: {summary.issue_source}\n")
            handle.write(f"Inputs: {len(summary.input_files)}\n")
            for index, item in enumerate(summary.input_files, 1):
                handle.write(f"  {index}. {item}\n")
            handle.write(f"Outputs: {len(summary.output_files)}\n")
            for index, item in enumerate(summary.output_files, 1):
                handle.write(f"  {index}. {item}\n")
            if summary.warnings:
                handle.write("\nWarnings:\n")
                for item in summary.warnings:
                    handle.write(f"  - {item}\n")
            if summary.errors:
                handle.write("\nItems needing review:\n")
                for item in summary.errors:
                    handle.write(f"  - {item}\n")
            if self.stamp_audit:
                handle.write("\nStamp verification:\n")
                for entry in self.stamp_audit:
                    handle.write(
                        f"  - {entry['type']} page {entry['page']}: {entry['text']} | "
                        f"verified={entry.get('verified')} | mode={entry.get('mode')} | "
                        f"reason={entry.get('reason', '')}\n"
                    )
        return log_path
